from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from pathlib import Path

OUT = Path(__file__).resolve().parents[1] / "output" / "AI_Editorial_Factory_Documento_Agentico.docx"

NAVY = "0B2545"
BLUE = "1677FF"
CYAN = "19C3E6"
LIGHT = "EAF2F8"
PALE = "F4F7FA"
INK = "17212B"
MUTED = "586879"
WHITE = "FFFFFF"
GREEN = "157347"
AMBER = "9A6700"

def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcPr.append(shd)
    shd.set(qn("w:fill"), fill)

def borders(cell, color="D6DEE6", size="6"):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:" + edge
        el = tcBorders.find(qn(tag))
        if el is None:
            el = OxmlElement(tag)
            tcBorders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), size)
        el.set(qn("w:color"), color)

def cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tcMar.find(qn("w:" + m))
        if node is None:
            node = OxmlElement("w:" + m)
            tcMar.append(node)
        node.set(qn("w:w"), str(v)); node.set(qn("w:type"), "dxa")

def set_cell_width(cell, dxa):
    tcPr = cell._tc.get_or_add_tcPr()
    tcW = tcPr.find(qn("w:tcW"))
    if tcW is None:
        tcW = OxmlElement("w:tcW"); tcPr.append(tcW)
    tcW.set(qn("w:w"), str(dxa)); tcW.set(qn("w:type"), "dxa")

def set_table_geometry(table, widths):
    table.autofit = False
    tblPr = table._tbl.tblPr
    tblW = tblPr.find(qn("w:tblW"))
    if tblW is None: tblW = OxmlElement("w:tblW"); tblPr.append(tblW)
    tblW.set(qn("w:w"), str(sum(widths))); tblW.set(qn("w:type"), "dxa")
    tblInd = tblPr.find(qn("w:tblInd"))
    if tblInd is None: tblInd = OxmlElement("w:tblInd"); tblPr.append(tblInd)
    tblInd.set(qn("w:w"), "120"); tblInd.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid): grid.remove(child)
    for w in widths:
        gc = OxmlElement("w:gridCol"); gc.set(qn("w:w"), str(w)); grid.append(gc)
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            set_cell_width(cell, widths[i]); cell_margins(cell); borders(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

def keep_with_next(p):
    p.paragraph_format.keep_with_next = True

def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("AI Editorial Factory  •  ")
    run.font.name = "Aptos"; run.font.size = Pt(8); run.font.color.rgb = RGBColor.from_string(MUTED)
    fld = OxmlElement("w:fldSimple"); fld.set(qn("w:instr"), "PAGE")
    paragraph._p.append(fld)

doc = Document()
sec = doc.sections[0]
sec.page_width = Inches(8.5); sec.page_height = Inches(11)
sec.top_margin = Inches(0.82); sec.bottom_margin = Inches(0.78)
sec.left_margin = Inches(0.85); sec.right_margin = Inches(0.85)
sec.header_distance = Inches(0.35); sec.footer_distance = Inches(0.35)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Aptos"; normal.font.size = Pt(10.2); normal.font.color.rgb = RGBColor.from_string(INK)
normal.paragraph_format.space_after = Pt(6); normal.paragraph_format.line_spacing = 1.16
for name, size, color, before, after in [
    ("Heading 1", 17, NAVY, 16, 7), ("Heading 2", 13.5, BLUE, 11, 5), ("Heading 3", 11.5, NAVY, 8, 4)
]:
    st = styles[name]; st.font.name = "Aptos Display"; st.font.size = Pt(size); st.font.bold = True
    st.font.color.rgb = RGBColor.from_string(color); st.paragraph_format.space_before = Pt(before)
    st.paragraph_format.space_after = Pt(after); st.paragraph_format.keep_with_next = True

for name in ("List Bullet", "List Number"):
    st = styles[name]; st.font.name = "Aptos"; st.font.size = Pt(10.2)
    st.paragraph_format.left_indent = Inches(.34); st.paragraph_format.first_line_indent = Inches(-.18)
    st.paragraph_format.space_after = Pt(3); st.paragraph_format.line_spacing = 1.12

if "Callout" not in styles:
    call = styles.add_style("Callout", WD_STYLE_TYPE.PARAGRAPH)
else: call = styles["Callout"]
call.font.name = "Aptos"; call.font.size = Pt(10); call.font.color.rgb = RGBColor.from_string(NAVY)
call.paragraph_format.left_indent = Inches(.18); call.paragraph_format.right_indent = Inches(.18)
call.paragraph_format.space_before = Pt(6); call.paragraph_format.space_after = Pt(9)

header = sec.header.paragraphs[0]
header.text = "AI EDITORIAL FACTORY  /  DOCUMENTO AGENTICO"
header.style = styles["Normal"]; header.runs[0].font.size = Pt(8); header.runs[0].font.bold = True
header.runs[0].font.color.rgb = RGBColor.from_string(MUTED)
add_page_number(sec.footer.paragraphs[0])

# Cover
p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(72)
r = p.add_run("AI EDITORIAL\nFACTORY")
r.font.name = "Aptos Display"; r.font.size = Pt(34); r.font.bold = True; r.font.color.rgb = RGBColor.from_string(NAVY)
p.paragraph_format.space_after = Pt(8)
p = doc.add_paragraph(); r = p.add_run("DOCUMENTO AGENTICO DI PROGETTO")
r.font.name = "Aptos Display"; r.font.size = Pt(16); r.font.bold = True; r.font.color.rgb = RGBColor.from_string(BLUE)
p.paragraph_format.space_after = Pt(20)
p = doc.add_paragraph("Architettura, agenti, workflow, dati, sicurezza, pubblicazione e roadmap della redazione multi-agente per manuali tecnici.")
p.style = styles["Subtitle"]; p.runs[0].font.name = "Aptos"; p.runs[0].font.size = Pt(13); p.runs[0].font.color.rgb = RGBColor.from_string(MUTED)
p.paragraph_format.space_after = Pt(28)
t = doc.add_table(rows=4, cols=2); t.alignment = WD_TABLE_ALIGNMENT.LEFT
data = [("Proprietario", "Daniel Meloni"), ("Versione documento", "1.0"), ("Data", "26 agosto 2026"), ("Stato progetto", "Fase 7 di 8 completata; Fase 8 progettata")]
for i,(a,b) in enumerate(data):
    t.cell(i,0).text=a; t.cell(i,1).text=b; shade(t.cell(i,0), LIGHT)
    t.cell(i,0).paragraphs[0].runs[0].font.bold=True
set_table_geometry(t,[2200,7160])
p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(24)
r=p.add_run("PRINCIPIO GUIDA"); r.bold=True; r.font.color.rgb=RGBColor.from_string(CYAN); r.font.size=Pt(9)
p=doc.add_paragraph("Gli agenti propongono, gli esseri umani decidono. Il contenuto originale resta immutabile e ogni pubblicazione richiede approvazione esplicita.", style="Callout")
pPr=p._p.get_or_add_pPr(); shd=OxmlElement("w:shd"); shd.set(qn("w:fill"), PALE); pPr.append(shd)
doc.add_page_break()

def h1(text): return doc.add_heading(text, level=1)
def h2(text): return doc.add_heading(text, level=2)
def para(text, bold_start=None):
    p=doc.add_paragraph()
    if bold_start and text.startswith(bold_start):
        r=p.add_run(bold_start); r.bold=True; p.add_run(text[len(bold_start):])
    else: p.add_run(text)
    return p
def bullets(items):
    for x in items: doc.add_paragraph(x, style="List Bullet")
def numbered(items):
    for i, x in enumerate(items, 1):
        p = doc.add_paragraph(f"{i}.\t{x}")
        p.paragraph_format.left_indent = Inches(.34)
        p.paragraph_format.first_line_indent = Inches(-.22)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.12
def table(headers, rows, widths):
    tb=doc.add_table(rows=1, cols=len(headers)); tb.alignment=WD_TABLE_ALIGNMENT.LEFT
    for i,h in enumerate(headers):
        c=tb.cell(0,i); c.text=h; shade(c,NAVY)
        for r in c.paragraphs[0].runs: r.font.bold=True; r.font.color.rgb=RGBColor.from_string(WHITE); r.font.size=Pt(9)
    for row in rows:
        cells=tb.add_row().cells
        for i,val in enumerate(row):
            cells[i].text=str(val)
            for p in cells[i].paragraphs:
                p.paragraph_format.space_after=Pt(1)
                for r in p.runs: r.font.size=Pt(9)
    set_table_geometry(tb,widths)
    doc.add_paragraph().paragraph_format.space_after=Pt(1)
    return tb

h1("1. Scopo e perimetro")
para("AI Editorial Factory è una piattaforma editoriale multi-agente che trasforma materiali tecnici grezzi — capitoli Markdown, codice, appendici, immagini e fonti — in opere verificate, illustrate e pubblicabili. Il primo pilota è Dataform in Pratica – Volume 1, ma il dominio applicativo è generico e riutilizzabile.")
h2("Obiettivi di sistema")
bullets(["Ridurre il lavoro ripetitivo senza cedere il controllo editoriale all’AI.", "Rendere tracciabile ogni proposta, fonte, revisione, approvazione ed esportazione.", "Separare l’analisi deterministica dall’elaborazione generativa.", "Produrre libri, articoli, lezioni e asset visuali coerenti a partire da una sola base editoriale.", "Estendere il modello dal singolo volume alla gestione completa di collane."])
h2("Principi non negoziabili")
table(["Principio","Garanzia"],[("Originale immutabile","Ogni intervento crea una nuova versione confrontabile e ripristinabile."),("Human-in-the-loop","I workflow si sospendono sui gate; nessuna pubblicazione automatica."),("Fonti verificabili","Gli URL proposti vengono aperti e validati prima di essere mostrati."),("Tenant isolation","RLS su tutte le tabelle esposte e controlli di appartenenza lato server."),("Niente funzioni simulate","Le azioni non implementate sono disabilitate e dichiarate prossime."),("Sviluppo economico","Provider mock predefiniti per percorrere il flusso senza crediti AI.")],[2450,6910])

h1("2. Stato del progetto")
para("La piattaforma ha completato sette fasi su otto. Il percorso operativo copre autenticazione, progetto, ingestione, riconoscimento della struttura, audit tecnico, revisione umana, diagrammi, copertina e pubblicazione in Markdown, HTML, PDF ed EPUB. La Fase 8, dedicata alle collane, dispone già di schema dati e modelli di dominio ma non ancora di interfaccia e workflow.")
table(["Area","Stato","Nota"],[("Fondazioni e autenticazione","Operativa","Sessioni SSR, PKCE, rotte protette, recupero password."),("Database e sicurezza","Operativa","Migration, storage privato, RLS e test PostgreSQL reali."),("Agenti e workflow capitolo","Parziale","4 agenti capitolo operativi su 12 previsti; gate umano attivo."),("Visual e copertine","Operativa / parziale","Diagrammi e cover programmatici; provider visuali reali da integrare."),("Pubblicazione","Operativa","Markdown, HTML, PDF, EPUB, articolo e lezione."),("Collane editoriali","Progettata","17 tabelle e regole presenti; UI e 6 agenti ancora assenti.")],[2600,1900,4860])
para("Baseline dichiarata dal repository: lint e typecheck senza errori, 469 test su 30 file, build con 26 rotte e 22 controlli smoke HTTP.")

h1("3. Architettura applicativa")
para("La soluzione adotta un monolite modulare Next.js con App Router. La durabilità dei processi lunghi è delegata a workflow asincroni, evitando microservizi prematuri. Le dipendenze sono unidirezionali: app e componenti dipendono dal dominio; il dominio dipende dagli adapter infrastrutturali Supabase e AI.")
table(["Livello","Responsabilità","Regola"],[("app/ e components/","Rotte, Server Component, Server Action, interfaccia utente.","Non contiene logica infrastrutturale sensibile."),("lib/ dominio","Validazione, regole editoriali, contratti agenti, workflow, publishing.","Espone funzioni testabili e provider-agnostic."),("lib/supabase e lib/ai","Persistenza, sessioni, storage e provider AI.","Implementa interfacce; i segreti restano server-only.")],[2100,4300,2960])
h2("Stack tecnologico")
table(["Ambito","Tecnologia"],[("Web","Next.js 16.3, React 19.2, TypeScript 5.9 strict"),("UI","Tailwind CSS 4.3, primitive accessibili"),("Dati e identità","Supabase: PostgreSQL, Auth, Storage, SSR"),("Validazione","Zod 4.4"),("Workflow","Workflow SDK 4.8"),("Contenuti","Unified, Remark, Rehype, unpdf, fflate"),("Output","React PDF, Markdown, HTML, EPUB"),("Qualità","Vitest, Playwright, ESLint 9, PGlite")],[2800,6560])
h2("Client Supabase")
bullets(["Browser client: publishable key, nessun segreto nel bundle.", "Server client: sessione da cookie e RLS attiva per componenti, action e route handler.", "Proxy client: rinnova token e cookie e governa i redirect di autenticazione.", "Admin client: service role confinato al server e ammesso solo dopo verifica esplicita dell’organizzazione."])

h1("4. Modello agentico")
para("L’agente è un operatore specializzato con un contratto di input/output validato da Zod, configurazione di provider e modello, prompt versionato e tracciamento dell’esecuzione. Il runner non concede potere di pubblicazione: produce risultati strutturati che diventano proposte revisionabili.")
h2("Ciclo di esecuzione standard")
numbered(["Preparare il contesto minimo necessario: progetto, capitolo, versione, fonti e regole editoriali.", "Validare l’input con lo schema dell’agente e registrare l’esecuzione.", "Selezionare il provider tramite registry comune; in sviluppo usare il mock.", "Eseguire analisi deterministiche prima della generazione AI.", "Validare l’output strutturato; rifiutare risposte incomplete o fuori contratto.", "Persistire risultato, motivazioni, riferimenti e metadati di costo/modello.", "Sospendere il workflow e richiedere revisione o approvazione umana.", "Applicare soltanto gli elementi approvati creando una nuova versione."])
h2("Agenti di capitolo")
table(["Agente","Responsabilità","Stato"],[("Technical Accuracy","Verifica codice, configurazioni e affermazioni tecniche.","Operativo"),("Source & Claims","Rileva claim senza fonte e propone riferimenti verificabili.","Operativo"),("Structure & Flow","Valuta struttura, ordine e continuità didattica.","Operativo"),("Editorial Revision","Propone una nuova versione del testo senza alterare l’originale.","Operativo"),("Style & Tone","Uniformità di voce, leggibilità e convenzioni.","Previsto"),("Terminology","Glossario, acronimi e coerenza lessicale.","Previsto"),("Code Review","Correttezza e qualità degli esempi eseguibili.","Previsto"),("Learning Design","Obiettivi, prerequisiti, esercizi e progressione.","Previsto"),("Accessibility","Alt text, struttura semantica e fruibilità.","Previsto"),("Visual Planning","Individua diagrammi e illustrazioni utili.","Previsto"),("Metadata & SEO","Metadati per web, blog e discovery.","Previsto"),("Publication Readiness","Controllo finale dei criteri di pubblicazione.","Previsto")],[2400,5260,1700])
h2("Agenti di collana — Fase 8")
table(["Agente","Missione"],[("Series Architect","Struttura, posizionamento e roadmap della collana."),("Series Curriculum","Progressione didattica, prerequisiti, lacune e duplicazioni."),("Series Consistency","Coerenza editoriale, terminologica, visiva e tecnica fra volumi."),("Series Visual Director","Identità comune e variazioni controllate per volume."),("Cross-Volume Reference","Integrità dei riferimenti fra capitoli e volumi."),("Series Publishing","Catalogo, schede volume, scaffale e pacchetto promozionale.")],[3000,6360])
para("Regola di governance: tutti gli agenti propongono e nessuno applica modifiche definitive senza autorizzazione umana.", bold_start="Regola di governance:")

h1("5. Workflow editoriale end-to-end")
numbered(["Creazione del progetto: autore, titolo, lingua, volume e brief editoriale.", "Ingestione: caricamento ZIP in storage privato con protezioni contro path traversal e zip bomb.", "Classificazione: riconoscimento di front matter, parti, capitoli, appendici e asset; ordinamento numerico.", "Confronto strutturale: indice dichiarato contro cartelle e file realmente presenti.", "Audit: analisi di SQLX, SQL, JavaScript, claim, fonti, struttura e rischi editoriali.", "Ricerca fonti: catalogo ufficiale, biblioteca del progetto e ricerca web verificata.", "Proposta: nuova versione con modifiche granulari, motivazioni e riferimenti.", "Revisione: diff per righe e parole, accettazione parziale, editing manuale, commenti e ripristino.", "Visual: diagrammi deterministici e illustrazioni versionate con approvazione.", "Cover Studio: fronte, dorso e quarta; testi e logo composti programmaticamente.", "Preflight: controllo stato asset, approvazioni, struttura, fonti e configurazione output.", "Pubblicazione: generazione di export e accesso tramite URL firmato."])
h2("Stati e gate")
para("Ogni esecuzione passa attraverso stati condivisi fra TypeScript e PostgreSQL: in attesa, in esecuzione, sospesa per input umano, completata o fallita. Il gate spezza deliberatamente l’automazione tra proposta e applicazione. Rifiuto e approvazione parziale sono esiti normali, non eccezioni.")

h1("6. Fonti, ricerca e verificabilità")
para("Il sistema combina tre canali mantenendo visibile l’origine di ciascun risultato.")
table(["Canale","Meccanismo","Vincolo"],[("Catalogo ufficiale","Indice curato e aggiornabile dalle sitemap della documentazione.","Non propone una pagina generica quando non esiste una corrispondenza pertinente."),("Biblioteca progetto","Link e PDF caricati dall’utente; PDF indicizzati pagina per pagina.","Ogni risultato conserva origine e pagina."),("Web discovery","Provider di ricerca AI, seguito da apertura reale dell’URL.","Le pagine non raggiungibili non vengono mostrate.")],[2100,4300,2960])
bullets(["Ogni claim rilevante deve essere collegabile a una fonte o marcato come non verificato.", "Titolo e metadati della fonte derivano dalla pagina letta, non dall’immaginazione del modello.", "Le fonti proposte entrano in una coda di accettazione o scarto umano.", "La bibliografia può essere ricostruita dalle fonti effettivamente approvate."])

h1("7. Versionamento e revisione umana")
para("Il contenuto sorgente non viene sovrascritto. Ogni proposta produce una versione figlia collegata al capitolo e all’esecuzione che l’ha generata. La workbench di revisione consente confronto, approvazione per porzioni, modifica manuale, commenti e rollback.")
h2("Invarianti")
bullets(["Una proposta deve essere attribuibile ad agente, provider, modello, prompt e timestamp.", "L’approvazione deve registrare identità, data e perimetro degli elementi accettati.", "Un asset approvato può sostituire il precedente, che passa a superseded senza essere cancellato.", "La pubblicazione usa soltanto versioni e asset nello stato ammesso dal preflight.", "Il ripristino crea un nuovo evento/versione; non cancella la storia."])

h1("8. Sistema visuale e copertine")
para("Il progetto separa nettamente ciò che deve essere tecnicamente esatto da ciò che può essere generativo.")
table(["Famiglia","Produzione","Proprietà"],[("Diagrammi tecnici","Mermaid o SVG da codice","Deterministici, riproducibili, sempre accompagnati da alt text."),("Illustrazioni","Adapter visuale configurabile","Prompt, negative prompt, provider, modello, seed, costo e parentela versionati."),("Copertine","Composizione programmatica + eventuale sfondo generato","Titolo, autore, logo, dorso e geometria restano controllabili e verificabili."),("Anteprime corsi","SVG 16:9 da codice","Identità coerente, testo esatto e asset autosufficiente.")],[2100,3500,3760])
h2("Regole della copertina")
bullets(["Larghezza foglio = 2 × pagina + dorso + 2 × abbondanza.", "Il dorso può derivare da mm/pagina, pagine per pollice o valore fisso del fornitore.", "Il valore definitivo si blocca solo con il numero definitivo di pagine.", "Sotto 6 mm il testo verticale del dorso viene omesso.", "Il logo dello strumento viene composto senza ritaglio e non viene ridisegnato dall’AI.", "ISBN validato; il barcode non è composto sulla copertina, ma il generatore EAN-13 resta testato e disponibile."])

h1("9. Pubblicazione e derivazioni")
table(["Output","Caratteristiche"],[("Markdown","Base portabile e leggibile, coerente con la struttura approvata."),("HTML","Conversione tramite Unified/Remark/Rehype e sanitizzazione."),("PDF","Composizione con React PDF e stile editoriale controllato."),("EPUB","Pacchetto e-book con metadati e navigazione."),("Lezione","Derivazione didattica per corso, con anteprima dedicata."),("Articolo","Derivazione blog con struttura e metadati dedicati.")],[2200,7160])
para("Gli export sono persistiti come record di pubblicazione e serviti dal bucket privato tramite URL firmati. La generazione non equivale alla pubblicazione: il sistema richiede che il gate umano e il preflight siano soddisfatti.")

h1("10. Modello dati e sicurezza")
para("Lo schema Supabase copre identità e organizzazioni, progetti e fonti, struttura editoriale, agenti e workflow, revisioni, asset visuali, pubblicazioni e collane. Gli enum condividono il vocabolario operativo con il codice TypeScript.")
h2("Controlli di sicurezza")
bullets(["RLS abilitata e forzata sulle tabelle esposte, inclusa la futura area collane.", "Service role esclusivamente server-side e preceduta da verifica esplicita del tenant.", "Bucket privati e URL firmati a breve durata per sorgenti, asset ed export.", "Sanitizzazione HTML e neutralizzazione XML nei testi inseriti in SVG.", "Validazione Zod sugli ingressi e limiti su dimensioni, quantità e struttura degli archivi.", "Protezione dai redirect esterni e rinnovo sicuro della sessione via proxy.", "Audit trail per azioni privilegiate, workflow, approvazioni e pubblicazioni."])
h2("Minacce coperte dall’ingestione")
table(["Rischio","Difesa"],[("Path traversal","Normalizzazione e guardia dei percorsi prima dell’estrazione."),("Zip bomb","Limiti su file, dimensioni compresse/decompresse e volume totale."),("File inattesi","Classificazione, whitelist logica e manifest verificato."),("Cross-tenant access","RLS e filtri organizzativi lato server."),("Asset orfani","Rollback dello storage se il record database non viene creato.")],[2600,6760])

h1("11. Collane editoriali — target Fase 8")
para("La collana è un dominio autonomo. Il legame progetto-volume vive in series_volumes, unica fonte di verità, così da evitare duplicazioni e consentire metadati specifici come numero, ISBN, dipendenze, data prevista e deroghe.")
h2("Ereditarietà delle regole")
table(["Stato","Significato"],[("inherited","Il volume segue la regola di collana e ne riceve gli aggiornamenti."),("overridden","Il volume usa una variante locale con motivazione obbligatoria."),("locked","La regola è inderogabile; ogni tentativo viene rifiutato.")],[2300,7060])
para("Le versioni di stile pubblicate sono immutabili. Una modifica genera proposta, analisi d’impatto, anteprima per volume e approvazione. I volumi già pubblicati non vengono aggiornati automaticamente: l’esito è una proposta di nuova edizione.")
h2("Fondamenta già presenti")
bullets(["17 tabelle di collana con RLS, indici e vincoli.", "Modelli TypeScript e stati del volume.", "Risoluzione resolveRule() con rifiuto delle deroghe su regole locked.", "Versionamento di linea editoriale, contenuti condivisi e glossario.", "Modello per riferimenti incrociati, piani di rilascio e issue di coerenza."])
h2("Componenti ancora da implementare")
bullets(["Rotte e interfaccia /series.", "Sei agenti di collana e relativi contratti.", "Workflow multi-volume e gate di impatto.", "Vista Scaffale con i dorsi affiancati.", "Export di catalogo, style guide, glossario e pacchetto promozionale."])

h1("12. Qualità, test e osservabilità")
table(["Livello","Copertura"],[("Unit test","Schemi, agenti, routing AI, ingestione, fonti, cover, publishing, derivazioni e utilità."),("Database test","Migration reali e RLS su PostgreSQL via PGlite."),("Component test","Navigazione, accessibilità e interazioni principali."),("E2E","Autenticazione e percorsi protetti con Playwright."),("Smoke test","Controlli HTTP contro la build di produzione."),("Static quality","ESLint, TypeScript strict, Prettier e validazione ambiente.")],[2300,7060])
h2("Telemetry minima per esecuzione agente")
bullets(["ID di esecuzione e correlazione col workflow.", "Agente, versione prompt, provider e modello.", "Input/output validati, durata, tentativi, esito e motivo di errore.", "Token/costo quando disponibili, senza registrare segreti.", "Fonti consultate e decisioni del gate umano."])

h1("13. Configurazione e deployment")
para("Il deployment target è Vercel con Supabase gestito. Le variabili pubbliche sono separate dai segreti e la validazione è pigra: la build può completarsi senza credenziali, mentre la prima operazione reale restituisce un errore esplicito.")
table(["Gruppo","Variabili principali"],[("Applicazione","NEXT_PUBLIC_APP_URL"),("Supabase pubblico","NEXT_PUBLIC_SUPABASE_URL, NEXT_PUBLIC_SUPABASE_PUBLISHABLE_KEY"),("Supabase server","SUPABASE_SERVICE_ROLE_KEY"),("AI testo","AI_TEXT_PROVIDER, AI_TEXT_MODEL, chiave provider"),("AI immagini","AI_IMAGE_PROVIDER, AI_IMAGE_MODEL"),("AI ricerca","AI_SEARCH_PROVIDER, AI_SEARCH_MODEL")],[2400,6960])
h2("Runbook essenziale")
numbered(["Installare dipendenze e compilare .env.local da .env.example.", "Eseguire check:env senza stampare segreti.", "Applicare le migration Supabase e verificare bucket e policy RLS.", "Eseguire lint, typecheck, test e build.", "Eseguire lo smoke test sulla build di produzione.", "Configurare URL applicazione e callback Auth su Supabase e Vercel.", "Distribuire soltanto dopo approvazione esplicita; mantenere un piano di rollback."])

h1("14. Governance operativa degli agenti")
h2("Matrice RACI sintetica")
table(["Attività","Agente","Editor umano","Sistema"],[("Analizzare","R","A/C","Supporta"),("Proporre modifica","R","A","Versiona"),("Accettare/scartare","C","A/R","Registra"),("Applicare versione","—","A","R, dopo gate"),("Generare export","R","A","Esegue preflight"),("Pubblicare","—","A/R","Blocca senza requisiti")],[2400,1900,2100,2960])
h2("Politiche operative")
bullets(["Minimo privilegio: ogni agente riceve soltanto contesto e strumenti necessari.", "Fail closed: output non valido o fonte non verificabile non avanza il workflow.", "Idempotenza: uno step ripetuto non deve duplicare versioni, asset o pubblicazioni.", "Riproducibilità: prompt, modello, seed e dipendenze devono essere rintracciabili.", "Budget: limiti per progetto, esecuzione e provider; mock come fallback di sviluppo.", "Escalation: conflitti, bassa confidenza e impatto su contenuti pubblicati richiedono revisione umana."])

h1("15. Roadmap raccomandata")
table(["Priorità","Iniziativa","Criterio di uscita"],[("P0","Completare gli 8 agenti capitolo mancanti","Contratti, runner, UI, test e gate coerenti con i 4 esistenti."),("P0","Preflight di pubblicazione unificato","Nessun export finale con fonti, versioni o asset non approvati."),("P1","Provider visuali reali","Adapter, cost tracking, retry e moderazione, senza cambiare il dominio."),("P1","Fase 8: UI e workflow collane","Creazione collana, ereditarietà, impatto, agenti e vista Scaffale."),("P1","Osservabilità agentica","Dashboard di esecuzioni, errori, costi, tempi e approvazioni."),("P2","Policy di valutazione continua","Dataset di casi, regression test e scorecard per agente."),("P2","Disaster recovery","Backup, ripristino verificato e runbook operativo.")],[1100,3650,4610])

h1("16. Criteri di accettazione globali")
bullets(["Un progetto completo attraversa ingestione, audit, revisione, visual, cover e pubblicazione senza modificare l’originale.", "Ogni proposta agentica è strutturata, attribuibile, versionata e reversibile.", "Ogni passaggio irreversibile o pubblico richiede approvazione umana esplicita.", "Fonti e URL sono verificabili; l’assenza di evidenza viene dichiarata.", "RLS impedisce accessi fra organizzazioni e nessun segreto raggiunge il client.", "Gli output pubblicati includono soltanto contenuti e asset approvati.", "Diagrammi tecnici sono deterministici; illustrazioni e copertine sono riproducibili e accessibili.", "Lint, typecheck, test database/unit/E2E, build e smoke risultano verdi.", "La collana applica regole ereditate, traccia deroghe e protegge i volumi pubblicati."])

h1("17. Glossario")
table(["Termine","Definizione"],[("Agente","Componente specializzato che produce un output strutturato e revisionabile."),("Workflow","Processo durevole composto da step, stati, retry e gate."),("Gate umano","Punto di sospensione che richiede una decisione esplicita."),("Claim","Affermazione tecnica verificabile che dovrebbe avere una fonte."),("Versione","Snapshot immutabile collegato al contenuto precedente e alla sua origine."),("Asset","Diagramma, immagine, copertina o file pubblicabile con stato e metadati."),("Derivazione","Contenuto secondario — articolo o lezione — generato dal materiale approvato."),("Preflight","Controllo finale dei prerequisiti di pubblicazione."),("RLS","Row Level Security: isolamento dei dati applicato dal database."),("Superseded","Stato di una versione sostituita ma conservata nello storico.")],[2200,7160])

h1("18. Fonti interne del documento")
para("Questo documento consolida lo stato rilevato nel repository al 26 agosto 2026. Le fonti principali sono:")
for src in ["README.md", "docs/architecture.md", "docs/agents.md", "docs/workflows.md", "docs/database.md", "docs/security.md", "docs/visual.md", "docs/publishing.md", "docs/deployment.md", "docs/series.md", "src/lib/agents/*", "src/workflows/*", "supabase/migrations/*", "tests/*"]:
    doc.add_paragraph(src, style="List Bullet")
para("Nota: le sezioni di roadmap e governance operativa distinguono esplicitamente le capacità già implementate dalle raccomandazioni per il completamento.")

# document properties
doc.core_properties.title = "AI Editorial Factory — Documento agentico di progetto"
doc.core_properties.subject = "Architettura, agenti, workflow e governance"
doc.core_properties.author = "Daniel Meloni"
doc.core_properties.keywords = "AI Editorial Factory, agenti, workflow, editoria, Next.js, Supabase"

OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUT)
print(OUT)
